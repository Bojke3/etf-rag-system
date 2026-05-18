from abc import ABC, abstractmethod
import logging
import os
from pathlib import Path
from typing import Dict, Type

from .document import Document
from .ocr import OCRHandler

logger = logging.getLogger(__name__)


class DocumentLoader(ABC):
    @abstractmethod
    def load(self, file_path: str) -> Document:
        pass

    @abstractmethod
    def validate(self, file_path: str) -> bool:
        pass


class PDFLoader(DocumentLoader):
    def load(self, file_path: str) -> Document:
        path = Path(file_path)
        text = ""
        ocr_used = False

        try:
            from pypdf import PdfReader

            text_parts = []
            with open(path, "rb") as f:
                reader = PdfReader(f)

                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)

            text = "\n".join(text_parts)
        except Exception as exc:
            logger.warning("PDF text extraction failed for %s: %s", path, exc)

        ocr_enabled = os.getenv("ETF_RAG_ENABLE_OCR", "1") != "0"

        if not text.strip() and ocr_enabled:
            ocr_text = OCRHandler().extract_text_from_pdf(str(path))
            if ocr_text.strip():
                text = ocr_text
                ocr_used = True

        return Document(
            text=text,
            metadata={
                "source": str(path),
                "file_name": path.name,
                "file_type": "pdf",
                "ocr_used": ocr_used,
            },
        )

    def validate(self, file_path: str) -> bool:
        path = Path(file_path)
        return path.exists() and path.is_file() and path.suffix.lower() == ".pdf"


class WordLoader(DocumentLoader):
    def load(self, file_path: str) -> Document:
        from docx import Document as DocxDocument

        path = Path(file_path)
        doc = DocxDocument(path)

        text_parts = [paragraph.text for paragraph in doc.paragraphs]

        return Document(
            text="\n".join(text_parts),
            metadata={
                "source": str(path),
                "file_name": path.name,
                "file_type": "docx",
            },
        )

    def validate(self, file_path: str) -> bool:
        try:
            path = Path(file_path)

            if not path.exists() or path.suffix.lower() != ".docx":
                return False

            from docx import Document as DocxDocument

            DocxDocument(path)
            return True
        except Exception:
            return False


class TextLoader(DocumentLoader):
    def load(self, file_path: str) -> Document:
        path = Path(file_path)

        with open(path, "r", encoding="utf-8") as f:
            text = f.read()

        return Document(
            text=text,
            metadata={
                "source": str(path),
                "file_name": path.name,
                "file_type": "txt",
            },
        )

    def validate(self, file_path: str) -> bool:
        try:
            path = Path(file_path)

            if not path.exists() or path.suffix.lower() != ".txt":
                return False

            with open(path, "r", encoding="utf-8") as f:
                f.read(100)

            return True
        except Exception:
            return False


class DocumentLoaderFactory:
    _loaders: Dict[str, Type[DocumentLoader]] = {
        ".pdf": PDFLoader,
        ".docx": WordLoader,
        ".txt": TextLoader,
    }

    @classmethod
    def get_loader(cls, file_path: str) -> DocumentLoader:
        suffix = Path(file_path).suffix.lower()

        if suffix not in cls._loaders:
            raise ValueError(f"Unsupported file format: {suffix}")

        return cls._loaders[suffix]()

    @classmethod
    def load_document(cls, file_path: str) -> Document:
        loader = cls.get_loader(file_path)

        if not loader.validate(file_path):
            raise ValueError(f"Invalid or unreadable file: {file_path}")

        return loader.load(file_path)
