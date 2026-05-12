"""Data layer - Document loading and preprocessing"""

from abc import ABC, abstractmethod
from pathlib import Path
from typing import List, Dict, Any
import logging

logger = logging.getLogger(__name__)

class DocumentLoader(ABC):
    """Abstract base class for document loaders"""
    
    @abstractmethod
    def load(self, file_path: str) -> str:
        """Load document and return text"""
        pass
    
    @abstractmethod
    def validate(self, file_path: str) -> bool:
        """Validate if file can be loaded"""
        pass

class PDFLoader(DocumentLoader):
    """Loader for PDF documents"""
    
    def load(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            from pypdf import PdfReader
            text = []
            with open(file_path, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text())
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            return ""
    
    def validate(self, file_path: str) -> bool:
        """Check if file is valid PDF"""
        try:
            from pypdf import PdfReader
            with open(file_path, 'rb') as f:
                PdfReader(f)
            return True
        except:
            return False

class WordLoader(DocumentLoader):
    """Loader for Word documents"""
    
    def load(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = [paragraph.text for paragraph in doc.paragraphs]
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error loading Word document {file_path}: {e}")
            return ""
    
    def validate(self, file_path: str) -> bool:
        """Check if file is valid Word document"""
        try:
            from docx import Document
            Document(file_path)
            return True
        except:
            return False

class TextLoader(DocumentLoader):
    """Loader for plain text documents"""
    
    def load(self, file_path: str) -> str:
        """Read text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            return ""
    
    def validate(self, file_path: str) -> bool:
        """Check if file can be read as text"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                f.read(100)
            return True
        except:
            return False

class DocumentLoaderFactory:
    """Factory for creating appropriate document loaders"""
    
    _loaders = {
        '.pdf': PDFLoader(),
        '.docx': WordLoader(),
        '.doc': WordLoader(),
        '.txt': TextLoader(),
    }
    
    @classmethod
    def get_loader(cls, file_path: str) -> DocumentLoader:
        """Get appropriate loader based on file extension"""
        suffix = Path(file_path).suffix.lower()
        if suffix not in cls._loaders:
            raise ValueError(f"Unsupported file format: {suffix}")
        return cls._loaders[suffix]
    
    @classmethod
    def load_document(cls, file_path: str) -> str:
        """Load document using appropriate loader"""
        loader = cls.get_loader(file_path)
        if not loader.validate(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        return loader.load(file_path)
